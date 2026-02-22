# app.py
import streamlit as st
import requests
import io
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from docx import Document

OLLAMA_API_DEFAULT = "http://localhost:11434/api/generate"
MODEL_NAME_DEFAULT = "granite4:micro"
DEFAULT_USERNAME = "Director"

st.set_page_config(page_title="Coffee with Cinema", layout="wide")
#bg
page_bg = """
<style>
[data-testid="stAppViewContainer"] {
    background-image: url("https://images.unsplash.com/photo-1524985069026-dd778a71c7b4"); /* Cinematic film reel */
    background-size: cover;
    background-position: center;
    background-repeat: no-repeat;
}
[data-testid="stSidebar"] {
    background-color: rgba(0,0,0,0.8); /* darker sidebar */
}
</style>
"""
st.markdown(page_bg, unsafe_allow_html=True)
st.title("Coffee with Cinema")
st.caption("An AI studio for storytellers, directors, and dreamers.")

#side
with st.sidebar:
    st.header("Settings")
    username = st.text_input("Your name", value=st.session_state.get("username", DEFAULT_USERNAME))
    st.session_state["username"] = username or DEFAULT_USERNAME
    st.markdown("---")
    st.write("Model endpoint (change only if needed)")
    st.session_state["ollama_api"] = st.text_input("Ollama API URL", value=st.session_state.get("ollama_api", OLLAMA_API_DEFAULT))
    st.session_state["model_name"] = st.text_input("Model name", value=st.session_state.get("model_name", MODEL_NAME_DEFAULT))
    st.markdown("---")
    st.write("Quick demo")
    if st.button("Load demo story"):
        st.session_state["storyline"] = (
            "A young detective in 1940s noir Los Angeles discovers a conspiracy that ties "
            "the city’s elite to a string of disappearances. As he digs deeper, he must choose "
            "between exposing the truth and protecting the woman he loves."
        )

# ---------- STORY INPUT ----------
st.subheader(f"Welcome , {st.session_state.get('username', DEFAULT_USERNAME)}")
storyline = st.text_area("Enter your story idea", value=st.session_state.get("storyline", ""), height=300)

col1, col2 = st.columns([1, 3])
with col1:
    if st.button("Generate Content"):
        if not storyline.strip():
            st.warning("Please enter a storyline.")
        else:
            st.session_state["storyline"] = storyline
            with st.spinner("Generating screenplay, characters, and sound design..."):
                sp_prompt = f"""You are a professional screenwriter. Based on the following story idea, write a detailed screenplay with proper formatting including scene headings, action lines, and dialogue.

Story Idea: {storyline}

Write a complete screenplay with:
- Scene headings (INT./EXT. LOCATION - TIME)
- Action descriptions
- Character dialogue
- Scene transitions

Screenplay:"""

                characters_prompt = f"""You are a character development expert. Based on the following story, create detailed character profiles.

Story: {storyline}

For each main character, provide:
- Name
- Age and Background
- Physical Description
- Personality Traits
- Character Arc
- Motivations
- Relationships

Character Profiles:"""

                sound_prompt = f"""You are a sound designer and composer. Based on the screenplay, create a detailed sound design and music plan for each scene.

Storyline: {storyline}

For each scene, describe:
- Background music style and mood
- Sound effects needed
- Ambient sounds
- Dialogue treatment
- Music cues and transitions

Sound Design Plan:"""

                def gen(prompt, max_tokens=1000):  # instead of 2000–3000
                    payload = {
                        "model": st.session_state.get("model_name", MODEL_NAME_DEFAULT),
                        "prompt": prompt,
                        "stream": False,
                        "options": {"temperature": 0.7, "num_predict": max_tokens}
                    }
                    api_url = st.session_state.get("ollama_api", OLLAMA_API_DEFAULT)
                    r = requests.post(api_url, json=payload, timeout=300)  # increase timeout
                    r.raise_for_status()
                    return r.json().get("response", "").strip()


                try:
                    screenplay = gen(sp_prompt, max_tokens=3000)
                    characters = gen(characters_prompt, max_tokens=2000)
                    sound_design = gen(sound_prompt, max_tokens=2000)
                    st.session_state["screenplay"] = screenplay
                    st.session_state["characters"] = characters
                    st.session_state["sound_design"] = sound_design
                    st.success("Generation complete.")
                except Exception as e:
                    st.error(f"Generation failed: {e}")

with col2:
    st.info("Tip: Use the demo story to show a fast live example during your pitch.")

# ---------- OUTPUTS ----------
if st.session_state.get("screenplay"):
    st.markdown("## Screenplay")
    st.code(st.session_state["screenplay"])
    st.download_button(
        "Download Screenplay TXT",
        st.session_state["screenplay"],
        file_name=f"screenplay_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
        mime="text/plain"
    )

    def make_pdf(text, title="Screenplay"):
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, alignment=TA_CENTER)
        body_style = ParagraphStyle('Body', parent=styles['BodyText'], fontSize=11, alignment=TA_JUSTIFY)
        story = [Paragraph(title, title_style), Spacer(1, 12)]
        for para in text.split("\n\n"):
            story.append(Paragraph(para.replace("\n", "<br/>"), body_style))
            story.append(Spacer(1, 6))
        doc.build(story)
        buf.seek(0)
        return buf

    pdf_buf = make_pdf(st.session_state["screenplay"])
    st.download_button(
        "Download Screenplay PDF",
        data=pdf_buf,
        file_name=f"screenplay_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
        mime="application/pdf"
    )

    def make_docx(text, title="Screenplay"):
        buf = io.BytesIO()
        doc = Document()
        doc.add_heading(title, level=1)
        for para in text.split("\n\n"):
            doc.add_paragraph(para)
        doc.save(buf)
        buf.seek(0)
        return buf
    
    docx_buf = make_docx(st.session_state["screenplay"])
    st.download_button(
         "Download Screenplay DOCX",
         data=docx_buf,
         file_name=f"screenplay_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ) 

if st.session_state.get("characters"):
         st.markdown("## Characters")
         st.text_area("Character Profiles", value=st.session_state["characters"], height=220)

if st.session_state.get("sound_design"): 
    st.markdown("## Sound Design")
    st.text_area("Sound Design Plan", value=st.session_state["sound_design"], height=220)

st.markdown("---")
st.caption("Keep Ollama running locally and ensure the model granite4:micro is pulled. Recommended RAM 16GB for smooth inference.")    
